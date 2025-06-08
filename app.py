import os 
import faiss
import pdfplumber
import requests
import numpy as np
import gradio as gr
from sentence_transformers import SentenceTransformer
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from langchain.docstore.document import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import docx  

GEMINI_API_KEY = "Enter your API_KEY "
os.environ["GOOGLE_API_KEY"] = GEMINI_API_KEY

embedder = SentenceTransformer('all-MiniLM-L6-v2')

genai = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash", 
    temperature=0.7,
    convert_system_message_to_human=True
)

pdf_chunks = []
chunk_embeddings = None
index = None

def extract_chunks_from_pdf(pdf_file):
    text = ""
    with pdfplumber.open(pdf_file.name) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return [text[i:i+500] for i in range(0, len(text), 500)]

def extract_chunks_from_txt(txt_file):
    with open(txt_file.name, 'r', encoding='utf-8') as file:
        text = file.read()
    return [text[i:i+500] for i in range(0, len(text), 500)]

def extract_chunks_from_docx(docx_file):
    doc = docx.Document(docx_file.name)
    text = ""
    for para in doc.paragraphs:
        text += para.text + "\n"
    return [text[i:i+500] for i in range(0, len(text), 500)]

def create_vector_store(chunks):
    global chunk_embeddings, index
    chunk_embeddings = embedder.encode(chunks)
    dim = chunk_embeddings.shape[1]
    index = faiss.IndexFlatL2(dim)
    index.add(np.array(chunk_embeddings))

def retrieve_relevant_chunks(question, top_k=3):
    q_embedding = embedder.encode([question])
    distances, indices = index.search(np.array(q_embedding), top_k)
    return [pdf_chunks[i] for i in indices[0]]

def ask_gemini_with_langchain(question, context_chunks):
    docs = [Document(page_content=chunk) for chunk in context_chunks]
    chain = load_qa_chain(genai, chain_type="stuff")
    result = chain.invoke({"input_documents": docs, "question": question})
    return result["output_text"]

def upload_file(file):
    global pdf_chunks
    file_ext = file.name.split('.')[-1].lower()
    
    if file_ext == 'pdf':
        pdf_chunks = extract_chunks_from_pdf(file)
    elif file_ext == 'txt':
        pdf_chunks = extract_chunks_from_txt(file)
    elif file_ext == 'docx':
        pdf_chunks = extract_chunks_from_docx(file)
    else:
        return "Unsupported file type. Please upload a PDF, TXT, or DOCX file."

    create_vector_store(pdf_chunks)
    return "File uploaded and indexed!"
    
    try:
        response = requests.post(
            "http://127.0.0.1:5001/upload", 
            json={"chunks": pdf_chunks}
        )
        if response.status_code == 200:
            return "File uploaded, indexed, and synced with DB!"
        else:
            return f"File uploaded but failed to sync with DB: {response.text}"
    except Exception as e:
        return f"Error syncing with DB server: {str(e)}"


def answer_question(question, model_choice):
    if not question or not pdf_chunks:
        return "Please upload a file and enter a question."
    
    relevant = retrieve_relevant_chunks(question)

    if model_choice == "Gemini":
        return ask_gemini_with_langchain(question, relevant)
    else:
        return "Only Gemini is currently supported."

with gr.Blocks() as demo:
    gr.Markdown("## 📄 PDF, TXT, DOCX QA with Gemini (via LangChain)")

    with gr.Row():
        file_input = gr.File(label="Upload PDF, TXT, or DOCX", file_types=[".pdf", ".txt", ".docx"])
        upload_btn = gr.Button("Process File")

    status = gr.Textbox(label="Status")

    with gr.Row():
        question_input = gr.Textbox(label="Ask a question")
        model_choice = gr.Radio(["Gemini"], label="Choose LLM", value="Gemini")
    
    answer_output = gr.Textbox(label="Answer", lines=6)

    upload_btn.click(upload_file, inputs=file_input, outputs=status)
    question_input.submit(answer_question, inputs=[question_input, model_choice], outputs=answer_output)

if __name__ == "__main__":
    demo.launch()
